"""
document_processor.py - Utility functions for processing uploaded documents
"""

import os
import logging
from typing import Dict, Any, List, Optional
import pdfplumber
from docx import Document
import markdown
import re
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text content from a PDF file

    Args:
        file_path: Path to the PDF file

    Returns:
        Extracted text as a string
    """
    try:
        text_content = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
        return "\n\n".join(text_content)
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
        return f"Error processing PDF: {str(e)}"

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text content from a DOCX file

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text as a string
    """
    try:
        doc = Document(file_path)
        text_content = []

        for para in doc.paragraphs:
            if para.text:
                text_content.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text:
                        row_text.append(cell.text)
                if row_text:
                    text_content.append(" | ".join(row_text))

        return "\n\n".join(text_content)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
        return f"Error processing DOCX: {str(e)}"

def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text content from a TXT file

    Args:
        file_path: Path to the TXT file

    Returns:
        Extracted text as a string
    """
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
            return file.read()
    except Exception as e:
        logger.error(f"Error extracting text from TXT {file_path}: {str(e)}")
        return f"Error processing TXT: {str(e)}"

def extract_text_from_markdown(file_path: str) -> str:
    """
    Extract text content from a Markdown file

    Args:
        file_path: Path to the Markdown file

    Returns:
        Extracted text as a string (with markdown formatting removed)
    """
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
            md_content = file.read()

        # Convert markdown to HTML
        html_content = markdown.markdown(md_content)

        # Remove HTML tags
        clean_text = re.sub(r'<[^>]*>', '', html_content)

        return clean_text
    except Exception as e:
        logger.error(f"Error extracting text from Markdown {file_path}: {str(e)}")
        return f"Error processing Markdown: {str(e)}"

def extract_document_text(file_path: str) -> Optional[str]:
    """
    Extract text from a document based on its file extension

    Args:
        file_path: Path to the document

    Returns:
        Extracted text as a string or None if the file type is not supported
    """
    file_extension = Path(file_path).suffix.lower()

    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    elif file_extension == '.txt':
        return extract_text_from_txt(file_path)
    elif file_extension == '.md':
        return extract_text_from_markdown(file_path)
    else:
        logger.warning(f"Unsupported file type: {file_extension}")
        return None

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """
    Split text into overlapping chunks for better search and context

    Args:
        text: The text to split
        chunk_size: Maximum size of each chunk in characters
        overlap: Number of overlapping characters between chunks

    Returns:
        List of text chunks
    """
    if not text:
        return []

    chunks = []

    # Split by paragraphs first
    paragraphs = text.split('\n\n')
    current_chunk = ""

    for paragraph in paragraphs:
        # If adding this paragraph exceeds chunk size, save current chunk and start new one
        if len(current_chunk) + len(paragraph) > chunk_size:
            if current_chunk:
                chunks.append(current_chunk.strip())

            # Start new chunk with overlap from previous chunk
            if len(current_chunk) > overlap:
                current_chunk = current_chunk[-overlap:] + "\n\n" + paragraph
            else:
                current_chunk = paragraph
        else:
            if current_chunk:
                current_chunk += "\n\n" + paragraph
            else:
                current_chunk = paragraph

    # Add the last chunk if it exists
    if current_chunk:
        chunks.append(current_chunk.strip())

    return chunks

def prepare_document_for_indexing(doc_info: Dict[str, Any], subject_name: str,
                                  file_path: str) -> List[Dict[str, Any]]:
    """
    Prepare document for indexing in Azure AI Search

    Args:
        doc_info: Document information dictionary
        subject_name: Name of the subject
        file_path: Path to the document file

    Returns:
        List of document chunks ready for indexing
    """
    try:
        # Extract text from document
        document_text = extract_document_text(file_path)

        if not document_text:
            logger.warning(f"No text could be extracted from {file_path}")
            return []

        # Split text into chunks
        text_chunks = chunk_text(document_text)

        # Create indexable documents
        documents = []
        for i, chunk in enumerate(text_chunks):
            # Create a unique ID for each chunk
            chunk_id = f"{doc_info['_id']}_{i}"

            # Ensure field names match exactly with the index schema
            document = {
                "id": chunk_id,
                "document_id": doc_info['_id'],
                "document_name": doc_info['filename'],
                "subject_id": doc_info.get('subject_id', ''),
                "subject_name": subject_name,
                "chunk_id": i,
                "content": chunk,
                "file_path": file_path
            }

            documents.append(document)

        logger.info(f"Prepared {len(documents)} document chunks for indexing from {file_path}")
        return documents

    except Exception as e:
        logger.error(f"Error preparing document for indexing: {str(e)}")
        return []